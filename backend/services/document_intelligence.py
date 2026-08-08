"""Azure Document Intelligence service for text extraction."""
from __future__ import annotations

import asyncio
import os
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

# Cap any single DI call so a stuck poller can't hold the request forever.
# vercel.json sets maxDuration=300 for the function. Default to 280s so DI
# times out with a clean error before Vercel hard-kills it silently.
# Override with DI_TIMEOUT_SECONDS env var (e.g. "8" for Hobby plan).
_DI_TIMEOUT_SECONDS = float(os.environ.get("DI_TIMEOUT_SECONDS", "280"))


# Phrases that mark a row as belonging to the technical ratings/spec block
# (rated pressure, design temperature, NACE/sour-service compliance,
# H2S/CO2/chloride/pH limits, etc.) rather than the actual equipment
# description. Once a continuation row matches one of these, ALL further
# continuation rows for that same item are also treated as spec-block text
# (the block always spans several wrapped rows together, and later rows in
# the block — e.g. a bare "Metallic: T-20/350..." line — don't repeat a
# marker keyword themselves).
_SPEC_BLOCK_MARKERS = (
    "rated working pressure",
    "design temperature",
    "sour service",
    "partial pressure",
    "chloride concentration",
    "elemental sulphur",
    "manufactured in accordance",
    "nace mr0175",
    "iso 15156",
    "ph (min)",
    "temperature for metallic",
    "non metallic sealing material",
    "non-metallic sealing material",
)


def _is_spec_block_text(text: str) -> bool:
    t = text.lower()
    return any(marker in t for marker in _SPEC_BLOCK_MARKERS)


def _cell_span_bounds(cell) -> tuple[int, int] | None:
    """Cell's (start, end) offset range in DI's shared reading-order content."""
    spans = getattr(cell, "spans", None) or []
    if not spans:
        return None
    return min(s.offset for s in spans), max(s.offset + s.length for s in spans)


def _cell_geometry(cell) -> tuple[int, float, float, float, float] | None:
    """Cell's (page_number, min_x, max_x, min_y, max_y) from its bounding region."""
    regions = getattr(cell, "bounding_regions", None) or []
    if not regions or not regions[0].polygon:
        return None
    poly = regions[0].polygon
    xs, ys = poly[0::2], poly[1::2]
    return regions[0].page_number, min(xs), max(xs), min(ys), max(ys)


def _line_geometry(line) -> tuple[float, float, float, float] | None:
    """Line's (min_x, max_x, min_y, max_y) from its polygon."""
    poly = getattr(line, "polygon", None)
    if not poly:
        return None
    xs, ys = poly[0::2], poly[1::2]
    return min(xs), max(xs), min(ys), max(ys)


def _attach_geometric_descriptions(
    row_starts: list[int],
    row_geometry: dict[int, tuple[int, float, float]],
    table_x_range: tuple[float, float],
    pages,
    items: list[dict],
) -> None:
    """
    Fill in each item's description purely from bounding-region position,
    for tables where Document Intelligence recognized NO description column
    at all (the description text was segmented as free-standing page lines,
    not table cells) — the exact shape of the reported bug: a stray
    reference line between two rows was being matched to the wrong item by
    the LLM's free-text order guessing, because nothing structural
    constrained which line belonged to which row.

    Each row-start's own cells (ITEMS/QTY/PART NO.) give a reliable vertical
    band on the page (top Y to the next row-start's top Y, or its own
    bottom Y for the last row). Any page line whose top Y falls in that band
    AND whose horizontal position does not overlap the table's own column
    area belongs to that row — first such line is the primary description,
    any further lines are continuations — with zero text/keyword matching.
    """
    windows: list[tuple[int, float, float] | None] = []
    for pos, row_idx in enumerate(row_starts):
        geom = row_geometry.get(row_idx)
        if geom is None:
            windows.append(None)
            continue
        page_no, top_y, bottom_y = geom
        next_geom = row_geometry.get(row_starts[pos + 1]) if pos + 1 < len(row_starts) else None
        upper = next_geom[1] if next_geom else bottom_y
        windows.append((page_no, top_y, upper))

    table_x0, table_x1 = table_x_range
    for page in pages or []:
        for line in page.lines or []:
            geom = _line_geometry(line)
            if geom is None:
                continue
            x0, x1, y0, _y1 = geom
            if x0 <= table_x1 and table_x0 <= x1:
                continue  # overlaps the table's own column area, not free description text
            val = (line.content or "").strip()
            if not val:
                continue
            for pos, window in enumerate(windows):
                if window is None or window[0] != page.page_number:
                    continue
                _, top, upper = window
                if top <= y0 < upper:
                    item = items[pos]
                    if item["_spec_mode"] or _is_spec_block_text(val):
                        item["_spec_mode"] = True
                        item["specifications"] = ((item["specifications"] or "") + " " + val).strip()
                    elif not item["knownDescription"]:
                        item["knownDescription"] = val
                    else:
                        item["knownDescription"] = (item["knownDescription"] + " " + val).strip()
                    break

    for item in items:
        if not item["knownDescription"]:
            item["needsReview"] = True


def _build_confirmed_line_items(tables, pages=None) -> list[dict]:
    """
    Deterministically group Document Intelligence table rows into line items.

    Uses each table's own row/column structure — not free text — to decide
    row boundaries: a row is a CONTINUATION of a row-start when its
    index-like columns (an "ITEMS"/index column and/or a "QTY" column) are
    blank, regardless of what appears in other columns (e.g. a secondary
    part/reference number printed in the PART NO. column slot). This removes
    the ambiguity that otherwise forces the LLM to guess where one BOM/parts
    row ends and the next begins — the exact mistake that caused a single
    row to be split into multiple fabricated line items.

    Row-starts and their continuations are ordered/attached using each
    cell's ``spans`` offset into DI's shared reading-order content, not the
    raw ``row_index`` integer. The table-structure model that assigns
    row_index is a separate model from the one that determines reading
    order, and on scanned/borderless layouts it can mislabel which row_index
    a wrapped/unlabeled cell belongs to (e.g. a stray lot/heat reference with
    no ITEMS/QTY of its own) — attaching by span position instead of
    row_index self-corrects that mislabeling, since span offsets always
    reflect true document position, without any text/keyword heuristics.

    Column roles are inferred generically from header text ("description",
    "qty"/"quantity", "part", "item") so this works across different CoC
    table layouts, not just one specific template. Tables with no identifiable
    part-number column (e.g. a plain customer/PO/SO key-value table) are
    skipped.
    """
    results: list[dict] = []

    for table in tables or []:
        cells_by_row: dict[int, dict[int, str]] = {}
        header_by_col: dict[int, str] = {}
        row_span_bounds: dict[int, tuple[int, int]] = {}
        for cell in table.cells or []:
            content = (cell.content or "").replace("\n", " ").strip()
            cells_by_row.setdefault(cell.row_index, {})[cell.column_index] = content
            if getattr(cell, "kind", None) == "columnHeader" or cell.row_index == 0:
                existing = header_by_col.get(cell.column_index, "")
                header_by_col[cell.column_index] = (existing + " " + content).strip()
            bounds = _cell_span_bounds(cell)
            if bounds is not None:
                existing_bounds = row_span_bounds.get(cell.row_index)
                row_span_bounds[cell.row_index] = (
                    (min(existing_bounds[0], bounds[0]), max(existing_bounds[1], bounds[1]))
                    if existing_bounds is not None else bounds
                )

        roles: dict[int, str] = {}
        for col, header in header_by_col.items():
            h = header.lower()
            if "description" in h:
                roles[col] = "description"
            elif "qty" in h or "quantity" in h:
                roles[col] = "qty"
            elif "part" in h:
                roles[col] = "partno"
            elif "item" in h:
                roles[col] = "items_index"
            elif "serial" in h:
                roles[col] = "serial"
            elif "invoice" in h:
                roles[col] = "invoice"
            elif "work order" in h or "w.o." in h or "w/o" in h:
                roles[col] = "workorder"
            else:
                # "lot" and "batch" are treated as ONE concept (lot/batch is a
                # single unified idea, not two) so a plain "Lot / Batch No."
                # column isn't miscounted as 2 concepts and wrongly flagged
                # as a combined column.
                has_lotbatch = "lot" in h or "batch" in h
                has_exp = "exp" in h
                has_so = "s/o" in h or "sales order" in h
                if sum((has_lotbatch, has_exp, has_so)) >= 2:
                    # Header smashes together 2+ distinct concepts (e.g.
                    # "S/O / LOT & BATCH / EXP.") — this is ONE combined
                    # column whose value must be preserved verbatim, not
                    # split apart.
                    roles[col] = "combined"
                elif has_lotbatch:
                    roles[col] = "lotbatch"
                elif has_exp:
                    roles[col] = "expiration"

        partno_col = next((c for c, r in roles.items() if r == "partno"), None)
        if partno_col is None:
            continue  # not a BOM/parts table

        index_cols = [c for c, r in roles.items() if r in ("items_index", "qty")]
        desc_col = next((c for c, r in roles.items() if r == "description"), None)
        qty_col = next((c for c, r in roles.items() if r == "qty"), None)
        serial_col = next((c for c, r in roles.items() if r == "serial"), None)
        lotbatch_col = next((c for c, r in roles.items() if r == "lotbatch"), None)
        expiration_col = next((c for c, r in roles.items() if r == "expiration"), None)
        combined_col = next((c for c, r in roles.items() if r == "combined"), None)
        invoice_col = next((c for c, r in roles.items() if r == "invoice"), None)
        workorder_col = next((c for c, r in roles.items() if r == "workorder"), None)

        def is_row_start(row: dict[int, str]) -> bool:
            if index_cols:
                return any(row.get(c, "").strip() for c in index_cols)
            return bool(row.get(partno_col, "").strip())

        def row_position(idx: int) -> int:
            # Rows with no span info (rare) keep their row_index position
            # instead of being reordered blindly.
            bounds = row_span_bounds.get(idx)
            return bounds[0] if bounds is not None else idx

        sorted_row_indices = sorted((i for i in cells_by_row.keys() if i != 0), key=row_position)

        # Table headers can wrap onto more than one physical row (e.g. "Rev."
        # on row 0 and "No." continuing on row 1, or "(if applicable)" split
        # across two rows). Treat every row before the first row that truly
        # looks like a data row (has content in an index/qty/partno column)
        # as part of the header block and skip it, instead of only skipping
        # row_index 0 — otherwise a leading header-continuation row gets
        # mistaken for a real (blank) line item.
        first_data_pos = next(
            (pos for pos, idx in enumerate(sorted_row_indices) if is_row_start(cells_by_row[idx])),
            None,
        )
        if first_data_pos is None:
            continue  # no recognizable data rows in this table

        # Fully blank rows (spacer/section-divider rows within the table)
        # carry no information — drop them before grouping instead of
        # letting them start a bogus item.
        data_rows = [
            idx for idx in sorted_row_indices[first_data_pos:]
            if any(v.strip() for v in cells_by_row[idx].values())
        ]
        row_starts = [idx for idx in data_rows if is_row_start(cells_by_row[idx])]
        row_start_set = set(row_starts)

        for pos, row_idx in enumerate(row_starts):
            row = cells_by_row[row_idx]
            qty_raw = row.get(qty_col, "").strip() if qty_col is not None else ""
            try:
                qty_val = int(qty_raw) if qty_raw else None
            except ValueError:
                qty_val = None
            item = {
                "partNumber": row.get(partno_col, "").strip() or None,
                "qty": qty_val,
                "knownDescription": row.get(desc_col, "").strip() if desc_col is not None else "",
                "specifications": None,
                "_spec_mode": False,
                "serial": None,
                "lotBatch": None,
                "expiration": None,
                "combined": None,
                "invoice": None,
                "workOrder": None,
                "needsReview": False,
            }
            # Capture serial/lot-batch/expiration/combined columns already
            # present on this SAME starting row (most BOM rows are a
            # single physical row with no continuation — without this,
            # that sibling-column data was silently dropped whenever
            # there was no wrapped continuation row to carry it via the
            # loop below). Columns that don't match any known role
            # (e.g. "Rev. No.", "PR Level") are intentionally ignored —
            # they aren't serial/lot/expiration data and would only add
            # noise for the LLM to misclassify.
            if serial_col is not None and row.get(serial_col, "").strip():
                item["serial"] = row[serial_col].strip()
            if lotbatch_col is not None and row.get(lotbatch_col, "").strip():
                item["lotBatch"] = row[lotbatch_col].strip()
            if expiration_col is not None and row.get(expiration_col, "").strip():
                item["expiration"] = row[expiration_col].strip()
            if combined_col is not None and row.get(combined_col, "").strip():
                item["combined"] = row[combined_col].strip()
            if invoice_col is not None and row.get(invoice_col, "").strip():
                item["invoice"] = row[invoice_col].strip()
            if workorder_col is not None and row.get(workorder_col, "").strip():
                item["workOrder"] = row[workorder_col].strip()

            # This item's span "block": every OTHER row (continuation, at ANY
            # row_index) whose own span falls inside this range belongs to
            # THIS item — regardless of what row_index DI assigned it. This
            # lets a wrapped cell that DI mislabeled with an out-of-sequence
            # row_index still attach to the correct item, instead of merging
            # into whichever item happened to be "current" in row_index order.
            lo = row_position(row_idx)
            hi = row_position(row_starts[pos + 1]) if pos + 1 < len(row_starts) else float("inf")

            for cont_idx in data_rows:
                if cont_idx in row_start_set or not (lo <= row_position(cont_idx) < hi):
                    continue
                cont_row = cells_by_row[cont_idx]
                for col, val in cont_row.items():
                    val = val.strip()
                    if not val:
                        continue
                    if col == desc_col:
                        if item["_spec_mode"] or _is_spec_block_text(val):
                            item["_spec_mode"] = True
                            item["specifications"] = ((item["specifications"] or "") + " " + val).strip()
                        else:
                            item["knownDescription"] = (item["knownDescription"] + " " + val).strip()
                    elif col == serial_col:
                        item["serial"] = ((item["serial"] or "") + " " + val).strip()
                    elif col == lotbatch_col:
                        item["lotBatch"] = ((item["lotBatch"] or "") + " " + val).strip()
                    elif col == expiration_col:
                        item["expiration"] = ((item["expiration"] or "") + " " + val).strip()
                    elif col == combined_col:
                        item["combined"] = ((item["combined"] or "") + " " + val).strip()
                    elif col == invoice_col:
                        item["invoice"] = ((item["invoice"] or "") + " " + val).strip()
                    elif col == workorder_col:
                        item["workOrder"] = ((item["workOrder"] or "") + " " + val).strip()
                    elif col == partno_col and serial_col is None and combined_col is None:
                        # Some templates print a secondary reference/serial
                        # number in the PART NO. column slot on a wrapped
                        # continuation row (no dedicated serial column at
                        # all) — treat it as this item's serial rather than
                        # dropping it.
                        item["serial"] = ((item["serial"] or "") + " " + val).strip()
                    # else: ignore other unclassified columns (Rev, PR Level,
                    # etc.) — noise, not part of this data model.

            # A description column exists in this table, but nothing in this
            # item's span block carried description text — a structural gap
            # DI itself couldn't resolve. Flag it instead of letting
            # free-text matching guess and risk attaching the wrong nearby
            # line (the exact failure this replaces).
            if desc_col is not None and not item["knownDescription"]:
                item["needsReview"] = True

            results.append(item)

        # This table has NO description column at all in its own cell
        # structure (desc_col is None) — reconstruct descriptions from
        # bounding-region geometry instead of leaving every item MISSING for
        # the LLM to guess from a flattened, unlinked text stream.
        if desc_col is None and pages:
            row_geometry: dict[int, tuple[int, float, float]] = {}
            xs_all: list[float] = []
            for cell in table.cells or []:
                geom = _cell_geometry(cell)
                if geom is None:
                    continue
                page_no, x0, x1, y0, y1 = geom
                xs_all.extend((x0, x1))
                existing = row_geometry.get(cell.row_index)
                row_geometry[cell.row_index] = (
                    (page_no, min(existing[1], y0), max(existing[2], y1))
                    if existing is not None else (page_no, y0, y1)
                )
            if xs_all and row_starts:
                table_x_range = (min(xs_all), max(xs_all))
                _attach_geometric_descriptions(
                    row_starts, row_geometry, table_x_range, pages, results[-len(row_starts):]
                )

    return results


def _render_confirmed_items_block(items: list[dict]) -> str:
    if not items:
        return ""

    lines = [
        "---CONFIRMED LINE ITEMS (grouped by deterministic layout/table "
        "analysis — these partNumber/qty values and row groupings, AND any "
        "serial/lotBatch/expiration/soLotBatchExp/invoice/workOrder/"
        "specifications values shown below, are AUTHORITATIVE; do not "
        "re-split, merge, re-derive, or change them. Only fill in "
        "'description' where marked MISSING, using the plain document text "
        "below in the same top-to-bottom order. Where marked NOT FOUND IN "
        "DOCUMENT STRUCTURE, leave description null instead — do not search "
        "free text for it. Never fold a "
        "'specifications' value into 'description' — they are separate "
        "fields.)---"
    ]
    for i, item in enumerate(items, start=1):
        if item.get("needsReview"):
            desc = (
                "(NOT FOUND IN DOCUMENT STRUCTURE — leave description null; "
                "do NOT guess it from surrounding text)"
            )
        else:
            desc = item["knownDescription"] or "(MISSING — find matching text in document)"
        parts = [f"Item {i}: partNumber={item['partNumber']} | qty={item['qty']} | description={desc}"]
        if item.get("serial"):
            parts.append(f"serial={item['serial']}")
        if item.get("lotBatch"):
            parts.append(f"lotBatch={item['lotBatch']}")
        if item.get("expiration"):
            parts.append(f"expiration={item['expiration']}")
        if item.get("combined"):
            parts.append(f"soLotBatchExp={item['combined']}")
        if item.get("invoice"):
            parts.append(f"invoice={item['invoice']}")
        if item.get("workOrder"):
            parts.append(f"workOrder={item['workOrder']}")
        if item.get("specifications"):
            parts.append(f"specifications={item['specifications']}")
        lines.append(" | ".join(parts))
    return "\n".join(lines)


async def extract_text(file_bytes: bytes, filename: str) -> tuple[str, bool]:
    """
    Extract raw text from a document using Azure Document Intelligence.

    Uses the **async** Document Intelligence SDK so the long-poll loop
    runs as real awaitable I/O — no thread pool, no blocked event loop.

    Returns:
        (extracted_text, is_ocr_needed)
        is_ocr_needed is True when DI returned very little or no text
        (e.g. scanned/image-only PDF).
    """
    from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
    from azure.core.credentials import AzureKeyCredential

    # Let KeyError propagate — missing env vars are a configuration error, not
    # a recoverable extraction failure.  The caller (ingest pipeline) will
    # catch this and surface it as a proper error message to the user.
    endpoint = os.environ["DOCUMENT_INTELLIGENCE_ENDPOINT"]
    key = os.environ["DOCUMENT_INTELLIGENCE_KEY"]
    model_id = os.environ.get("DI_MODEL_ID", "prebuilt-layout")
    # Hard cap at 30 pages per file — configurable via DI_MAX_PAGES env var.
    max_pages = int(os.environ.get("DI_MAX_PAGES", "30"))

    try:
        async with DocumentIntelligenceClient(
            endpoint=endpoint,
            credential=AzureKeyCredential(key),
        ) as client:
            poller = await client.begin_analyze_document(
                model_id=model_id,
                body=file_bytes,
                content_type="application/octet-stream",
                pages=f"1-{max_pages}",
            )
            result = await asyncio.wait_for(poller.result(), timeout=_DI_TIMEOUT_SECONDS)

        lines: list[str] = []
        for page in result.pages or []:
            for line in page.lines or []:
                lines.append(line.content)

        text = "\n".join(lines).strip()

        confirmed_items = _build_confirmed_line_items(result.tables, result.pages)
        confirmed_block = _render_confirmed_items_block(confirmed_items)
        if confirmed_block:
            text = text + "\n\n" + confirmed_block

        is_ocr_needed = len(text) < 100

        logger.info("DI extracted %d chars from %s", len(text), filename)
        return text, is_ocr_needed

    except asyncio.TimeoutError:
        logger.error("DI timed out after %.0fs for %s", _DI_TIMEOUT_SECONDS, filename)
        raise RuntimeError(
            f"Document Intelligence timed out after {_DI_TIMEOUT_SECONDS:.0f}s — "
            "set DI_TIMEOUT_SECONDS env var or check your Azure DI resource."
        )
    except Exception as exc:
        logger.warning("DI extraction failed for %s: %s", filename, exc)
        exc_str = str(exc)
        if "InvalidContentLength" in exc_str or "too large" in exc_str:
            raise RuntimeError(
                "The file size is too large for Azure Document Intelligence. "
                "If your Azure resource is on the Free (F0) tier, the maximum file size is 4 MB. "
                "Please upgrade your Azure resource to the Standard (S0) tier to support files up to 500 MB, "
                "or compress your document to reduce its file size below 4 MB."
            )
        # Re-raise API/credential errors so the caller can surface them.
        # Only swallow the error for a genuine "no text" response.
        raise RuntimeError(f"Document Intelligence extraction failed: {exc}") from exc


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from a Word document.

    Dispatches by magic bytes:
      - ``PK\\x03\\x04``        →  modern .docx (ZIP/Open XML)  → python-docx
      - ``\\xD0\\xCF\\x11\\xE0...`` →  legacy .doc (OLE2 binary)    → olefile parser
    """
    _DOCX_MAGIC = b"PK\x03\x04"
    _OLE2_MAGIC = b"\xD0\xCF\x11\xE0\xA1\xB1\x1A\xE1"

    if not file_bytes:
        return ""

    if file_bytes[:4] == _DOCX_MAGIC:
        return _extract_docx_text(file_bytes)
    if file_bytes[:8] == _OLE2_MAGIC:
        return _extract_doc_ole_text(file_bytes)

    # Unknown — try both, return whichever yields more text
    a = _extract_docx_text(file_bytes)
    b = _extract_doc_ole_text(file_bytes)
    return a if len(a) >= len(b) else b


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from a modern .docx (Open XML / ZIP) file."""
    import io
    from docx import Document  # type: ignore

    try:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    paragraphs.append("  |  ".join(cells))
        return "\n".join(paragraphs)
    except Exception as exc:
        logger.warning("python-docx extraction failed: %s", exc)
        return ""


def _extract_doc_ole_text(file_bytes: bytes) -> str:
    """
    Extract text from a legacy binary .doc (Word 97-2003, OLE2 Compound File).

    Implements the MS-DOC Word Binary File Format piece-table walker:
      1. Open OLE2 container, read FIB from start of ``WordDocument`` stream.
      2. Read ``fWhichTblStm`` (FibBase.flags bit 9) to pick ``0Table`` vs ``1Table``.
      3. Read ``fcClx`` / ``lcbClx`` from ``FibRgFcLcb97`` (offsets 0x01A2 / 0x01A6).
      4. Locate ``PlcPcd`` inside the ``Clx`` block.
      5. Walk N+1 character-position entries and N piece descriptors.  For each
         piece, ``FcCompressed`` bit 30 (``fCompressed``) decides 1-byte CP1252
         (offset // 2) or 2-byte UTF-16 LE encoding.
    """
    import io
    import struct

    try:
        import olefile  # type: ignore
    except ImportError:
        logger.warning("olefile not installed — cannot read legacy .doc")
        return ""

    try:
        ole = olefile.OleFileIO(io.BytesIO(file_bytes))
        if not ole.exists("WordDocument"):
            logger.warning(".doc has no WordDocument stream")
            ole.close()
            return ""

        wd = ole.openstream("WordDocument").read()
        if len(wd) < 0x1A6 + 4:
            logger.warning(".doc WordDocument stream too short")
            ole.close()
            return ""

        # FibBase.flags is a 16-bit word at offset 0x000A
        flags = struct.unpack_from("<H", wd, 0x000A)[0]
        f_which_table = (flags >> 9) & 1
        table_name = "1Table" if f_which_table else "0Table"

        if not ole.exists(table_name):
            logger.warning(".doc missing %s stream", table_name)
            ole.close()
            return ""
        table = ole.openstream(table_name).read()
        ole.close()

        # FibRgFcLcb97: fcClx at 0x01A2, lcbClx at 0x01A6
        fc_clx = struct.unpack_from("<I", wd, 0x01A2)[0]
        lcb_clx = struct.unpack_from("<I", wd, 0x01A6)[0]

        if lcb_clx == 0 or fc_clx + lcb_clx > len(table):
            logger.warning(".doc Clx pointer out of range")
            return ""

        clx = table[fc_clx : fc_clx + lcb_clx]

        # Walk Clx: each entry is either Prc (type=0x01) or Pcdt (type=0x02).
        # We want the Pcdt, which contains the PlcPcd piece table.
        i = 0
        plc_pcd: bytes = b""
        while i < len(clx):
            entry_type = clx[i]
            if entry_type == 0x01:
                # Prc: 1-byte type, 2-byte cbGrpprl, then cbGrpprl bytes
                if i + 3 > len(clx):
                    break
                cb = struct.unpack_from("<h", clx, i + 1)[0]
                i += 3 + max(cb, 0)
            elif entry_type == 0x02:
                # Pcdt: 1-byte type, 4-byte lcb, then PlcPcd
                if i + 5 > len(clx):
                    break
                lcb = struct.unpack_from("<I", clx, i + 1)[0]
                plc_pcd = clx[i + 5 : i + 5 + lcb]
                break
            else:
                break

        if not plc_pcd:
            logger.warning(".doc PlcPcd not found")
            return ""

        # PlcPcd = (N+1) CPs (each 4 bytes) + N PCDs (each 8 bytes)
        # 4*(N+1) + 8*N = len(plc_pcd)  =>  N = (len - 4) / 12
        n = (len(plc_pcd) - 4) // 12
        if n <= 0:
            return ""

        cps = struct.unpack_from(f"<{n + 1}I", plc_pcd, 0)
        pcd_offset = 4 * (n + 1)

        parts: list[str] = []
        for k in range(n):
            cp_start = cps[k]
            cp_end = cps[k + 1]
            if cp_end <= cp_start:
                continue

            # PCD is 8 bytes; FcCompressed is the 4-byte dword at offset 2
            fc_compressed = struct.unpack_from("<I", plc_pcd, pcd_offset + 8 * k + 2)[0]
            f_compressed = (fc_compressed >> 30) & 1
            fc = fc_compressed & 0x3FFFFFFF

            char_count = cp_end - cp_start
            if f_compressed:
                # 1 byte / char, CP1252; offset is fc/2 in stream
                start = fc // 2
                raw = wd[start : start + char_count]
                try:
                    text = raw.decode("cp1252", errors="replace")
                except Exception:
                    text = raw.decode("latin-1", errors="replace")
            else:
                # 2 bytes / char, UTF-16 LE
                raw = wd[fc : fc + char_count * 2]
                try:
                    text = raw.decode("utf-16-le", errors="replace")
                except Exception:
                    text = ""

            parts.append(text)

        # Normalise Word control characters
        text = "".join(parts)
        text = (
            text.replace("\r", "\n")
                .replace("\x07", "\t")  # cell mark
                .replace("\x0C", "\n")  # page break
                .replace("\x0B", "\n")  # line break
                .replace("\x13", "")    # field begin
                .replace("\x14", "")    # field separator
                .replace("\x15", "")    # field end
                .replace("\x01", "")    # embedded object
                .replace("\x02", "")
                .replace("\x05", "")    # annotation ref
                .replace("\x08", "")
        )
        # Collapse runs of blank lines
        lines = [ln.strip() for ln in text.splitlines()]
        cleaned = [ln for ln in lines if ln]
        return "\n".join(cleaned)

    except Exception as exc:
        logger.warning("legacy .doc OLE parse failed: %s", exc)
        return ""
